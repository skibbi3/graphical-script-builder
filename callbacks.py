from dash import callback, Output, Input, State, dcc
import json
import base64
from docx import Document

@callback(
    Output('script-table', 'data', allow_duplicate=True),
    Input('add-row', 'n_clicks'),
    State('script-table', 'data'),
    State('script-table', 'columns'), 
    prevent_initial_call=True
)
def add_row(n_clicks, rows, columns):
    if n_clicks > 0:
        # Increments the scene number by 1
        previous_scene = int(rows[-1]['scene'])
        rows.append({
            'scene':str(previous_scene+1).zfill(3),
            'duration':0,
            'video':'',
            'audio':''})
    return rows

@callback(
    Output('script-table', 'data', allow_duplicate=True),
    Input('script-table', 'data'),
    prevent_initial_call=True
)
def update_table(rows):
    total_time = 0
    new_rows = []
    i = 0
    for row in rows:
        # Ensure the row numbers are consistent
        new = row
        new['scene'] = str(i).zfill(3)
        i+=1

        # Calculate duration
        if new['audio'] is not None:
            new['duration'] = len(new['audio'].split(" "))//2
            total_time += new['duration']

        # Add it all to the new row
        new_rows.append(new)

    return new_rows


@callback(
    Output('duration', 'value', allow_duplicate=True),
    Input('script-table', 'data'),
    prevent_initial_call=True
)
def update_duration(data):
    """
    Calculates the total duration of the script and returns it to the duration text field.
    """
    total_time = 0
    for row in data:
        if row['duration'] != '':
            total_time += int(row['duration'])

    minutes = total_time//60
    seconds = total_time - (minutes * 60)
    return str(minutes) + " minutes " + str(seconds) + " seconds"

@callback(
    Output('save-work-download', 'data', allow_duplicate=True),
    Input('save-work-button', 'n_clicks'),
    State('program-title', 'value'),
    State('program-code', 'value'),
    State('duration', 'value'),
    State('date', 'value'),
    State('version', 'value'),
    State('writtenby', 'value'),
    State('editedby', 'value'),
    State('script-table', 'data'),
    prevent_initial_call=True
)
def save_work(n_clicks, program_title, program_code, duration, date, version, writtenby, editedby, table_data):
    """
    Saves all the input fields, including the table, as a JSON file and returns it for the user the download to their machine.
    """
    contents = {
        "Title of Program": program_title,
        "Program Code": program_code,
        "Duration": duration,
        "Date": date,
        "Version": version,
        "Written By": writtenby,
        "Edited By": editedby,
        "Table": table_data
    }

    with open("./output/progress.json", "w") as outfile:
        json.dump(contents, outfile)

    return dcc.send_file("./output/progress.json")

@callback(
    Output('export-work-download', 'data', allow_duplicate=True),
    Input('export-work-button', 'n_clicks'),
    State('program-title', 'value'),
    State('program-code', 'value'),
    State('duration', 'value'),
    State('date', 'value'),
    State('version', 'value'),
    State('writtenby', 'value'),
    State('editedby', 'value'),
    State('script-table', 'data'),
    prevent_initial_call=True
)
def export_work(n_clicks, program_title, program_code, duration, date, version, writtenby, editedby, table_data):
    document = Document()
    document.add_paragraph("Title of program: " + str(program_title))
    document.add_paragraph("Program code: " + str(program_code))
    document.add_paragraph("Duration: " + str(duration))
    document.add_paragraph("Date: " + str(date))
    document.add_paragraph("Version: " + str(version))
    document.add_paragraph("Written By: " + str(writtenby))
    document.add_paragraph("Edited By: " + str(editedby))

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "SCENE"
    hdr_cells[1].text = "DURATION"
    hdr_cells[2].text = "VISION"
    hdr_cells[3].text = "AUDIO"
    for row in table_data:
        row_cells = table.add_row().cells
        row_cells[0].text = row['scene']
        row_cells[1].text = str(row['duration'])
        row_cells[2].text = row['video']
        row_cells[3].text = row['audio']

    document.save("./output/script.docx")

    return dcc.send_file("./output/script.docx")

@callback(
    Output('program-title', 'value'),
    Output('program-code', 'value'),
    Output('duration', 'value'),
    Output('date', 'value'),
    Output('version', 'value'),
    Output('writtenby', 'value'),
    Output('editedby', 'value'),
    Output('script-table', 'data'),
    Input('previous-work-upload', 'contents'),
    prevent_initial_call = True
)
def upload_work(contents):
    content_type, content_string = contents.split(",")
    decoded = base64.b64decode(content_string)

    content_dict = json.loads(decoded)

    return content_dict['Title of Program'], content_dict['Program Code'], content_dict['Duration'], content_dict['Date'], content_dict['Version'], content_dict['Written By'], content_dict['Edited By'], content_dict['Table']